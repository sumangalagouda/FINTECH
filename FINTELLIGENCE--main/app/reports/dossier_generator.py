import io
from flask import jsonify, send_file, request
from docx import Document
import openpyxl
from app.models.case import Case
from app.models.transaction import Transaction
from app.reports.report_generator import reports_bp

@reports_bp.route('/dossier/<case_id>', methods=['GET'])
def generate_dossier(case_id):
    authority = request.args.get('authority', 'bank_fraud')
    case = Case.query.get(case_id)
    
    if not case:
        return jsonify({"error": "Case not found"}), 404
        
    transactions = Transaction.query.filter_by(case_id=case_id, is_failed=False).all()
    
    if authority == 'auditor':
        # Generate Excel
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Raw Data"
        headers = ["Date", "Amount", "Type", "Sender", "Receiver", "Description", "Balance After"]
        ws.append(headers)
        
        for t in transactions:
            ws.append([
                t.date.strftime('%Y-%m-%d') if t.date else "",
                t.amount,
                t.type,
                t.sender_account,
                t.receiver_account,
                t.description,
                t.balance_after
            ])
            
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"Dossier_Auditor_{case_id}.xlsx",
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        
    else:
        from app.ai.ollama_client import call_ollama
        
        system_prompt = (
            f"You are a forensic data analyst compiling a strictly factual Case Dossier for the {authority.replace('_', ' ').title()}. "
            "Write highly professional, investigative text based on the provided facts. "
            "Do NOT hallucinate. Do not include asterisks or markdown, just plain text paragraphs."
        )
        user_prompt = f"""
Please generate the intelligence analysis section of the CASE DOSSIER based on these case facts.

TEMPLATE TO FOLLOW:
PART I: EXECUTIVE SUMMARY
- Summary of Intelligence: [AI-generated paragraph summarizing the laundering or fraud mechanism]

PART II: FINANCIAL ANALYSIS & RED FLAGS
- Heatmap / Detector Breakdown: [Describe why the score is high based on the risk level]
- Recommended next steps for {authority.replace('_', ' ').title()}.

CASE FACTS:
- Case ID: {case_id}
- Account: {case.bank_name} - {case.account_number}
- Suspicion Score: {case.suspicion_score}/100
- Risk Level: {case.risk_level}
- Total Transactions: {len(transactions)}
"""
        try:
            dossier_ai_text = call_ollama(system_prompt, user_prompt, max_tokens=1000)
        except Exception as e:
            dossier_ai_text = f"[Error generating AI intelligence report: {str(e)}]\n\n"

        doc = Document()
        doc.add_heading(f"FINTELLIGENCE Dossier: {authority.replace('_', ' ').title()}", 0)
        
        doc.add_heading("Case Details", level=1)
        doc.add_paragraph(f"Case ID: {case_id}")
        doc.add_paragraph(f"Risk Score: {case.suspicion_score}")
        doc.add_paragraph(f"Account: {case.bank_name} - {case.account_number}")
        
        doc.add_heading("AI Intelligence & Analysis", level=1)
        for line in dossier_ai_text.split('\n'):
            if line.strip():
                clean_line = line.replace('**', '').replace('##', '').replace('*', '-')
                doc.add_paragraph(clean_line)
                
        doc.add_heading("Transaction Details Summary", level=1)
        doc.add_paragraph(f"Total Transactions analyzed: {len(transactions)}")
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"Dossier_{authority}_{case_id}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
